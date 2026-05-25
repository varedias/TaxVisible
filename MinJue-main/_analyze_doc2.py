from docx import Document

doc = Document(r'F:\Development\Java\IDEA_Projects\MinJue\软著用户手册.docx')

# Print all paragraphs with their style info
for i, p in enumerate(doc.paragraphs):
    style = p.style.name if p.style else 'None'
    text = p.text.strip()
    if text:
        print(f'[{i}] Style: {style} | {text[:150]}')

print('--- END ---')
print(f'Total paragraphs: {len(doc.paragraphs)}')

# Print table info
print(f'Total tables: {len(doc.tables)}')
for ti, t in enumerate(doc.tables):
    print(f'Table {ti}: {len(t.rows)} rows x {len(t.columns)} cols')
    for ri, row in enumerate(t.rows):
        if ri < 3:
            cells = [c.text.strip()[:30] for c in row.cells]
            print(f'  Row {ri}: {cells}')

# Print sections info
for si, section in enumerate(doc.sections):
    print(f'Section {si}: width={section.page_width}, height={section.page_height}')
    header = section.header
    if header and header.paragraphs:
        for hp in header.paragraphs:
            if hp.text.strip():
                print(f'  Header: {hp.text.strip()[:100]}')
    footer = section.footer
    if footer and footer.paragraphs:
        for fp in footer.paragraphs:
            if fp.text.strip():
                print(f'  Footer: {fp.text.strip()[:100]}')
