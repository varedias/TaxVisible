"""
B2B工业设备宣传与电商平台 - 软件设计说明书 V1.0
生成符合CPCC规范的Word文档（.docx）
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ============================================================
# 工具函数
# ============================================================

def set_cell_shading(cell, color_hex):
    """设置单元格底色"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_paragraph_spacing(paragraph, before=0, after=0, line_spacing=1.5):
    """设置段落间距"""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line_spacing

def add_formatted_paragraph(doc, text, style='Normal', font_name='宋体',
                             font_size=12, bold=False, alignment=None,
                             space_before=0, space_after=6, line_spacing=1.5,
                             first_line_indent=None, color=None):
    """添加格式化段落"""
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    if alignment is not None:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    if first_line_indent:
        pf.first_line_indent = Cm(first_line_indent)
    return p

def add_body_text(doc, text, indent=True, bold=False):
    """添加正文段落（小四号宋体，1.5倍行距，首行缩进2字符）"""
    return add_formatted_paragraph(
        doc, text, font_name='宋体', font_size=12,
        bold=bold, space_before=0, space_after=6,
        line_spacing=1.5,
        first_line_indent=0.74 if indent else None
    )

def add_heading1(doc, text):
    """一级标题"""
    p = doc.add_paragraph(style='Heading 1')
    run = p.add_run(text)
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    pf = p.paragraph_format
    pf.space_before = Pt(24)
    pf.space_after = Pt(12)
    pf.line_spacing = 1.5
    return p

def add_heading2(doc, text):
    """二级标题"""
    p = doc.add_paragraph(style='Heading 2')
    run = p.add_run(text)
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    pf = p.paragraph_format
    pf.space_before = Pt(18)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.5
    return p

def add_heading3(doc, text):
    """三级标题"""
    p = doc.add_paragraph(style='Heading 3')
    run = p.add_run(text)
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.5
    return p

def add_table(doc, headers, rows, col_widths=None):
    """添加格式化表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(10.5)
        run.font.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, 'D9E2F3')

    # 数据行
    for ri, row_data in enumerate(rows):
        for ci, cell_text in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(10.5)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 设置列宽
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    doc.add_paragraph()  # 表后空行
    return table

def add_figure_caption(doc, caption):
    """添加图表编号与标题"""
    p = add_formatted_paragraph(
        doc, caption, font_name='宋体', font_size=10.5,
        bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=6, space_after=12
    )
    return p

# ============================================================
# 主文档生成
# ============================================================

doc = Document()

# ---- 页面设置 ----
for section in doc.sections:
    section.page_width = Cm(21)       # A4
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2)

# ---- 设置默认字体 ----
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ---- 页眉 ----
header = doc.sections[0].header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = hp.add_run('B2B工业设备宣传与电商平台 V1.0    设计说明书')
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(128, 128, 128)

# ---- 页脚（页码） ----
footer = doc.sections[0].footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
# 添加自动页码域
run1 = fp.add_run()
fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
run1._element.append(fldChar1)
run2 = fp.add_run()
instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
run2._element.append(instrText)
run3 = fp.add_run()
fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
run3._element.append(fldChar2)


# ============================
# 封面页
# ============================
for _ in range(6):
    doc.add_paragraph()

add_formatted_paragraph(
    doc, 'B2B工业设备宣传与电商平台',
    font_name='宋体', font_size=28, bold=True,
    alignment=WD_ALIGN_PARAGRAPH.CENTER,
    space_before=0, space_after=12
)

add_formatted_paragraph(
    doc, '（Version 1.0）',
    font_name='等线', font_size=18, bold=False,
    alignment=WD_ALIGN_PARAGRAPH.CENTER,
    space_before=0, space_after=36
)

for _ in range(2):
    doc.add_paragraph()

add_formatted_paragraph(
    doc, '软', font_name='华文行楷', font_size=26, bold=False,
    alignment=WD_ALIGN_PARAGRAPH.CENTER,
    space_before=0, space_after=0
)
add_formatted_paragraph(
    doc, '件', font_name='华文行楷', font_size=26, bold=False,
    alignment=WD_ALIGN_PARAGRAPH.CENTER,
    space_before=0, space_after=0
)
add_formatted_paragraph(
    doc, '设', font_name='华文行楷', font_size=26, bold=False,
    alignment=WD_ALIGN_PARAGRAPH.CENTER,
    space_before=0, space_after=0
)
add_formatted_paragraph(
    doc, '计', font_name='华文行楷', font_size=26, bold=False,
    alignment=WD_ALIGN_PARAGRAPH.CENTER,
    space_before=0, space_after=0
)
add_formatted_paragraph(
    doc, '说', font_name='华文行楷', font_size=26, bold=False,
    alignment=WD_ALIGN_PARAGRAPH.CENTER,
    space_before=0, space_after=0
)
add_formatted_paragraph(
    doc, '明', font_name='华文行楷', font_size=26, bold=False,
    alignment=WD_ALIGN_PARAGRAPH.CENTER,
    space_before=0, space_after=0
)
add_formatted_paragraph(
    doc, '书', font_name='华文行楷', font_size=26, bold=False,
    alignment=WD_ALIGN_PARAGRAPH.CENTER,
    space_before=0, space_after=36
)

for _ in range(4):
    doc.add_paragraph()

add_formatted_paragraph(
    doc, '二〇二六年二月',
    font_name='宋体', font_size=18, bold=True,
    alignment=WD_ALIGN_PARAGRAPH.CENTER,
    space_before=36, space_after=0
)

doc.add_page_break()

# ============================
# 前言
# ============================
add_formatted_paragraph(
    doc, '前          言',
    font_name='宋体', font_size=22, bold=True,
    alignment=WD_ALIGN_PARAGRAPH.CENTER,
    space_before=24, space_after=18
)

add_body_text(doc, '欢迎使用B2B工业设备宣传与电商平台，本文档是该软件系统的设计说明书，旨在全面、系统地描述本软件的架构设计、功能模块划分、核心业务流程、数据结构模型、接口规范以及异常处理机制等关键技术内容。本文档面向软件著作权审查人员，作为软件设计鉴别材料的重要组成部分。')

add_body_text(doc, 'B2B工业设备宣传与电商平台是一款面向工业设备行业的企业级综合服务系统，涵盖AI视觉检测、自动化设备、工业相机、光源镜头、测量仪器、工业机器人等六大品类的设备展示与交易、设备租赁、供应商入驻与审核、行业内容发现以及后台运营管理等核心业务场景。')

add_body_text(doc, '本文档中的所有技术描述、架构图示、流程说明均严格基于系统现有源代码的分析结果，真实、准确地反映软件的设计思想与实现细节。若对本文档有任何疑问或建议，请与我们联系，谨谢！')

doc.add_page_break()

# ============================
# 目录页（占位）
# ============================
add_formatted_paragraph(
    doc, '目     录',
    font_name='黑体', font_size=22, bold=True,
    alignment=WD_ALIGN_PARAGRAPH.CENTER,
    space_before=24, space_after=18
)

# 手工目录
toc_items = [
    ('前言', ''),
    ('文件修订记录', ''),
    ('第一章 软件概述', ''),
    ('    1. 软件背景与用途', ''),
    ('    2. 开发目标与技术特点', ''),
    ('    3. 运行环境与技术架构', ''),
    ('第二章 系统架构设计', ''),
    ('    1. 整体架构说明', ''),
    ('    2. 系统架构图', ''),
    ('第三章 功能模块设计', ''),
    ('    1. 功能结构概览', ''),
    ('    2. 系统认证模块', ''),
    ('    3. 商品管理模块', ''),
    ('    4. 供应商管理模块', ''),
    ('    5. 租赁管理模块', ''),
    ('    6. 内容管理模块', ''),
    ('    7. 订单与购物车模块', ''),
    ('    8. 用户交互模块', ''),
    ('    9. 管理后台模块', ''),
    ('第四章 核心算法与流程', ''),
    ('    1. 用户认证与登录流程', ''),
    ('    2. JWT请求认证过滤流程', ''),
    ('    3. 用户注册流程', ''),
    ('    4. 购物车与订单创建流程', ''),
    ('    5. 订单状态流转', ''),
    ('    6. 供应商审核流程', ''),
    ('第五章 数据结构设计', ''),
    ('    1. 核心数据模型', ''),
    ('    2. 主要数据表说明', ''),
    ('    3. 缓存数据结构', ''),
    ('    4. 数据传输对象', ''),
    ('第六章 接口设计', ''),
    ('    1. 接口总体说明', ''),
    ('    2. 接口模块划分', ''),
    ('    3. 系统认证接口', ''),
    ('    4. 商品与分类接口', ''),
    ('    5. 购物车与订单接口', ''),
    ('    6. 管理后台核心接口', ''),
    ('    7. 接口调用时序', ''),
    ('第七章 异常处理设计', ''),
    ('    1. 全局异常处理机制', ''),
    ('    2. 异常处理流程', ''),
    ('    3. 主要异常场景', ''),
    ('    4. 事务回滚策略', ''),
]

for item_text, _ in toc_items:
    is_chapter = item_text.startswith('第') or item_text in ('前言', '文件修订记录')
    fs = 12 if is_chapter else 11
    bd = is_chapter
    add_formatted_paragraph(
        doc, item_text, font_name='宋体', font_size=fs,
        bold=bd, space_before=2, space_after=2, line_spacing=1.5
    )

doc.add_page_break()

# ============================
# 文件修订记录
# ============================
add_heading1(doc, '文件修订记录')

add_table(doc,
    headers=['版本号', '生成日期', '作者', '修订内容'],
    rows=[
        ['V1.0', '2026-02-06', '开发团队', '初始版本'],
        ['', '', '', ''],
        ['', '', '', ''],
    ],
    col_widths=[3, 3.5, 3, 6.5]
)

doc.add_page_break()

# ============================================================
# 第一章 软件概述
# ============================================================
add_heading1(doc, '第一章 软件概述')

add_heading2(doc, '1. 软件背景与用途')

add_body_text(doc, 'B2B工业设备宣传与电商平台是一款面向工业设备行业的B2B综合服务型软件系统，旨在为工业视觉检测、自动化设备、工业相机、光源镜头、测量仪器、工业机器人等领域的供应商与采购方搭建高效的线上对接平台。')

add_body_text(doc, '该平台覆盖设备产品展示与销售、设备租赁（融资租赁与经营租赁）、供应商入驻与认证审核、行业内容发现（视频测评、技术教程、行业资讯等）、用户互动（评论、点赞、收藏、分享）以及后台运营管理等核心业务场景。')

add_body_text(doc, '本软件通过前后端分离架构，后端提供标准化的RESTful API接口，为前端或第三方系统提供统一的数据服务与业务逻辑处理能力。系统支持三种用户角色——普通采购用户（USER）、供应商（SUPPLIER）和管理员（ADMIN），覆盖从商品浏览、下单交易、供应商入驻审核、设备租赁浏览到运营数据统计的完整业务链路。')

add_heading2(doc, '2. 开发目标与技术特点')

add_body_text(doc, '基于对源代码的深入分析，本软件在设计与实现中体现了以下核心目标与技术特点：')

features = [
    ('模块化设计', '系统采用模块化架构，按业务领域划分为系统认证模块（system）、商品管理模块（product）、供应商管理模块（supplier）、租赁模块（leasing）、内容管理模块（content）、订单与购物车模块（order）、用户交互模块（interaction）以及管理后台模块（admin），各模块职责明确、低耦合高内聚。'),
    ('安全认证体系', '基于Spring Security与JWT实现无状态令牌认证，结合图形验证码和邮箱验证码实现多因素身份验证，支持USER、SUPPLIER、ADMIN三种角色的差异化访问控制。JWT令牌采用HS256签名算法，有效期24小时。'),
    ('高性能缓存策略', '采用Redis实现购物车数据的持久化存储（Hash结构，30天自动过期），以及验证码的短时效缓存（5分钟），有效降低数据库压力并提升系统响应速度。Redis采用Lettuce客户端，通过JSON序列化器处理值对象。'),
    ('统一响应封装', '所有接口响应通过统一的Result<T>泛型包装类返回，包含状态码（code）、消息（message）和数据（data）三个标准字段，并配套全局异常处理机制，确保接口返回格式的一致性。'),
    ('事务一致性保障', '在订单创建、供应商审核等涉及多表操作的关键业务中，采用Spring声明式事务管理（@Transactional），确保数据操作的原子性和一致性。'),
    ('API文档自动化', '集成Knife4j（基于OpenAPI 3.0）自动生成在线API文档，便于前后端协作与接口调试。'),
]

for title, desc in features:
    p = doc.add_paragraph(style='Normal')
    run_title = p.add_run(f'（{features.index((title, desc)) + 1}）{title}：')
    run_title.font.name = '宋体'
    run_title._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run_title.font.size = Pt(12)
    run_title.font.bold = True
    run_desc = p.add_run(desc)
    run_desc.font.name = '宋体'
    run_desc._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run_desc.font.size = Pt(12)
    pf = p.paragraph_format
    pf.space_before = Pt(3)
    pf.space_after = Pt(3)
    pf.line_spacing = 1.5
    pf.first_line_indent = Cm(0.74)

add_heading2(doc, '3. 运行环境与技术架构')

add_body_text(doc, '本软件的运行环境与核心技术栈信息如下表所示：', bold=True)

# 硬件要求
add_body_text(doc, '硬件要求', bold=True, indent=False)
add_table(doc,
    headers=['类别', '基本要求'],
    rows=[
        ['服务器端', '内存4G以上；硬盘空间不低于40G；CPU双核及以上'],
        ['客户端', '支持现代浏览器的终端设备；内存2G及以上'],
    ],
    col_widths=[4, 12]
)

# 软件要求
add_body_text(doc, '软件环境要求', bold=True, indent=False)
add_table(doc,
    headers=['类别', '名称', '版本/要求'],
    rows=[
        ['开发语言', 'Java', 'JDK 17'],
        ['核心框架', 'Spring Boot', '3.2.2'],
        ['安全框架', 'Spring Security', '含JWT认证过滤器'],
        ['持久层框架', 'MyBatis-Plus', '3.5.5'],
        ['数据库', 'MySQL', 'InnoDB引擎，utf8mb4字符集'],
        ['缓存中间件', 'Redis', 'Lettuce客户端'],
        ['邮件服务', 'Spring Boot Starter Mail', 'SMTP协议'],
        ['接口文档', 'Knife4j', '4.5.0（OpenAPI 3.0）'],
        ['工具库', 'Hutool', '5.8.25'],
        ['JWT库', 'JJWT', '0.11.5（HS256算法）'],
        ['构建工具', 'Apache Maven', '-'],
        ['连接池', 'HikariCP', '内置'],
    ],
    col_widths=[3, 5, 8]
)

add_body_text(doc, '系统以Spring Boot可执行JAR方式独立部署运行，默认监听端口8999，通过CORS跨域过滤器支持前后端分离架构下的跨域请求调用。数据库连接池采用HikariCP，最大连接数配置为20个，连接超时时间30秒。Redis连接池最大活跃连接数为8个。')

doc.add_page_break()

# ============================================================
# 第二章 系统架构设计
# ============================================================
add_heading1(doc, '第二章 系统架构设计')

add_heading2(doc, '1. 整体架构说明')

add_body_text(doc, '本系统采用经典的分层架构模式，由上至下依次为客户端请求层、安全过滤层、控制器层（Controller）、服务层（Service）、数据访问层（Mapper/DAO）以及数据存储层（MySQL数据库与Redis缓存）。此外，系统还依赖外部SMTP邮件服务完成邮箱验证码的投递。')

add_body_text(doc, '客户端请求首先经过跨域过滤器（CorsFilter），获取跨域访问支持。随后由JWT认证过滤器（JwtAuthenticationTokenFilter）解析请求头中的Bearer Token，从数据库加载用户角色信息并注入Spring Security上下文。最终由Spring Security安全链（SecurityFilterChain）决定请求的放行或拦截，路由至对应的控制器处理。')

add_body_text(doc, '控制器层按业务领域划分为八大模块组，分别处理系统认证、商品管理、供应商管理、租赁管理、内容管理、订单与购物车、用户交互以及管理后台的请求。所有控制器通过统一的Result<T>包装类返回响应，异常情况由全局异常处理器（GlobalExceptionHandler）统一捕获处理。')

add_body_text(doc, '服务层承担核心业务逻辑运算，与数据访问层（MyBatis-Plus Mapper接口）交互完成数据持久化操作。其中，验证码服务（CaptchaService）和邮件服务（EmailService）通过Redis缓存实现验证码的临时存储；购物车服务（CartService）通过Redis Hash结构实现购物车数据的高性能读写；邮件服务额外依赖外部SMTP服务器完成验证码邮件的发送。')

add_heading2(doc, '2. 系统架构图')

add_body_text(doc, '下图展示了B2B工业设备宣传与电商平台的整体技术架构，清晰呈现从客户端到数据存储的完整分层关系和模块交互方式：', indent=False)

# 以文字表格形式呈现架构图
add_table(doc,
    headers=['层级', '组件名称', '职责说明'],
    rows=[
        ['客户端层', '前端应用 / 第三方系统', '发起HTTP/HTTPS请求，携带JWT令牌'],
        ['安全过滤层', 'CorsFilter', '处理跨域请求，允许所有来源访问'],
        ['安全过滤层', 'JwtAuthenticationTokenFilter', '解析Bearer Token，加载用户角色至SecurityContext'],
        ['安全过滤层', 'SecurityFilterChain', 'Spring Security安全链，无状态会话管理'],
        ['控制器层', 'SysUserController / CaptchaController / EmailController', '系统认证：登录注册、验证码、邮箱验证'],
        ['控制器层', 'PmsProductController / PmsCategoryController', '商品管理：商品列表、详情、分类'],
        ['控制器层', 'OmsSupplierController', '供应商管理：列表、详情、CRUD'],
        ['控制器层', 'OmsLeasingController', '租赁管理：设备列表、详情'],
        ['控制器层', 'CmsContentController', '内容管理：内容列表、详情、CRUD'],
        ['控制器层', 'OmsOrderController / CartController', '订单与购物车：下单、支付、购物车操作'],
        ['控制器层', 'AdminAuth / Dashboard / User / Product / Supplier / Order / Leasing / Interaction', '管理后台：全模块运营管理'],
        ['服务层', 'SysUserService', '用户认证、注册、密码重置、JWT生成'],
        ['服务层', 'CaptchaService', '图形验证码生成与校验（Redis缓存）'],
        ['服务层', 'EmailService', '邮箱验证码发送与校验（Redis + SMTP）'],
        ['服务层', 'PmsProductService / PmsCategoryService', '商品与分类业务逻辑'],
        ['服务层', 'OmsSupplierService', '供应商业务逻辑'],
        ['服务层', 'OmsLeasingService', '租赁设备业务逻辑'],
        ['服务层', 'CmsContentService', '内容业务逻辑（含浏览量统计）'],
        ['服务层', 'OmsOrderService', '订单创建、支付、取消（事务管理）'],
        ['服务层', 'CartService', '购物车操作（Redis Hash存储）'],
        ['服务层', 'CommentService / LikeService / FavoriteService / ShareService', '用户交互数据管理'],
        ['数据访问层', 'MyBatis-Plus Mapper接口', '数据库CRUD操作，自动映射驼峰命名'],
        ['基础设施', 'Result<T> / GlobalExceptionHandler / JwtUtil', '统一响应封装、全局异常处理、JWT工具'],
        ['基础设施', 'MybatisPlusConfig / RedisConfig', '分页插件、Redis序列化配置'],
        ['数据存储', 'MySQL（minjue_db）', '12张核心业务数据表'],
        ['数据存储', 'Redis', '验证码缓存、购物车Hash存储'],
        ['外部依赖', 'SMTP邮件服务', '邮箱验证码HTML邮件发送'],
    ],
    col_widths=[2.8, 5.5, 7.7]
)

add_figure_caption(doc, '图2-1 系统架构层级表')

doc.add_page_break()

# ============================================================
# 第三章 功能模块设计
# ============================================================
add_heading1(doc, '第三章 功能模块设计')

add_heading2(doc, '1. 功能结构概览')

add_body_text(doc, '系统功能按业务领域划分为八大功能模块：系统认证模块、商品管理模块、供应商管理模块、租赁管理模块、内容管理模块、订单与购物车模块、用户交互模块以及管理后台模块。各模块职责明确，通过服务层实现业务逻辑解耦。')

add_table(doc,
    headers=['模块编号', '模块名称', '主要功能', '核心组件'],
    rows=[
        ['M1', '系统认证模块', '用户注册、登录、密码重置、验证码、邮箱验证', 'SysUserController, CaptchaService, EmailService'],
        ['M2', '商品管理模块', '商品列表查询、详情、分类管理', 'PmsProductController, PmsCategoryController'],
        ['M3', '供应商管理模块', '供应商列表、详情、CRUD', 'OmsSupplierController'],
        ['M4', '租赁管理模块', '租赁设备列表、详情（融资/经营）', 'OmsLeasingController'],
        ['M5', '内容管理模块', '内容列表、详情、CRUD（文章/视频/Vlog）', 'CmsContentController'],
        ['M6', '订单与购物车模块', '购物车管理、购物车下单、直接下单、支付、取消', 'OmsOrderController, CartController'],
        ['M7', '用户交互模块', '评论、点赞、收藏、分享数据管理', 'CommentService, LikeService等'],
        ['M8', '管理后台模块', '全模块CRUD运营、仪表盘统计、审核、批量操作', 'Admin系列Controller（8个）'],
    ],
    col_widths=[2, 3.5, 5.5, 5]
)

add_figure_caption(doc, '图3-1 系统功能模块总览表')

add_heading2(doc, '2. 系统认证模块')

add_body_text(doc, '系统认证模块负责用户身份的注册、登录、密码重置以及认证辅助功能。该模块由SysUserController、CaptchaController和EmailController三个控制器协同工作，依赖SysUserService、CaptchaService和EmailService三个服务类实现核心业务逻辑。')

add_heading3(doc, '2.1 用户注册')
add_body_text(doc, '接收用户名、密码、确认密码、邮箱、昵称、角色（USER/SUPPLIER）等信息，经图形验证码校验、邮箱验证码校验和密码一致性检查后，对密码进行BCrypt加密，创建用户记录并持久化至数据库。用户名需全局唯一，若重复则抛出业务异常。注册成功后用户状态默认为启用（status=1），角色默认为USER。')

add_heading3(doc, '2.2 用户登录')
add_body_text(doc, '接收用户名、密码、图形验证码和登录角色参数，校验验证码有效性后，通过BCrypt密码比对和账号状态检查完成身份验证。系统实现了角色校验逻辑：管理员角色（ADMIN）可跨角色登录；非管理员角色需严格匹配——supplier角色需数据库角色为SUPPLIER，buyer角色需数据库角色为USER。全部校验通过后，使用JJWT库以username为Subject、HS256算法生成有效期为24小时的JWT令牌。')

add_heading3(doc, '2.3 密码重置')
add_body_text(doc, '通过图形验证码和邮箱验证码双重验证后，根据用户名和邮箱匹配用户记录，使用BCrypt对新密码加密后更新数据库中的密码字段和更新时间。')

add_heading3(doc, '2.4 图形验证码服务')
add_body_text(doc, '基于Hutool工具库的LineCaptcha生成200×80像素、4位字符、50条干扰线的图形验证码。以UUID为键存入Redis（键格式：captcha:{uuid}，5分钟有效期），返回UUID和Base64编码的图片数据。验证时采用一次性消费模式——校验完成后立即从Redis中删除该记录，验证码比较忽略大小写。')

add_heading3(doc, '2.5 邮箱验证码服务')
add_body_text(doc, '通过Hutool的RandomUtil生成6位随机数字验证码，以email:code:{type}:{email}为键存入Redis（5分钟有效期），其中type支持register、reset、login三种场景。通过JavaMailSender以SMTP协议发送HTML格式的验证码邮件，邮件主题为"【民崛智能】验证码"。验证后立即从Redis中删除记录。')

add_heading2(doc, '3. 商品管理模块')

add_body_text(doc, '商品管理模块提供面向用户端的商品浏览与搜索功能。PmsProductController支持分页查询商品列表，可按商品名称模糊搜索、按分类ID筛选，并支持五种排序策略：')

sort_items = [
    '按销量降序排列（sort=sales）',
    '按价格升序排列（sort=price-low）',
    '按价格降序排列（sort=price-high）',
    '按最新上架时间排列（sort=newest）',
    '综合排序——默认按销量和浏览量降序排列',
]
for item in sort_items:
    p = doc.add_paragraph(style='Normal')
    run = p.add_run(f'• {item}')
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(12)
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.5
    pf.left_indent = Cm(1.5)

add_body_text(doc, '商品列表默认仅展示上架状态（status=1）的商品，可通过includeOffShelf参数包含下架商品。商品详情接口在返回数据时自动递增浏览量字段（views+1）。PmsCategoryController提供商品分类的完整CRUD功能，分类支持通过parentId字段构建层级结构（0表示顶级分类）。')

add_heading2(doc, '4. 供应商管理模块')

add_body_text(doc, '供应商模块通过OmsSupplierController提供供应商的分页列表查询、详情查看以及创建、更新、删除等操作。供应商实体存储企业名称（name）、Logo图片URL（logo）、企业简介（description）、联系方式（contactInfo，JSON格式存储联系人姓名、电话、邮箱、地址）、认证状态（isVerified：0-待审核、1-已认证、2-审核拒绝）和关联用户ID（userId）等信息。供应商通过userId与系统用户表建立关联关系。')

add_heading2(doc, '5. 租赁管理模块')

add_body_text(doc, '租赁模块面向用户端提供设备租赁信息的浏览功能。OmsLeasingController支持按租赁类型（financing融资租赁/operating经营租赁）、设备名称进行筛选查询，默认仅展示上架状态的设备，按已租次数（leased）降序排列。')

add_body_text(doc, '融资租赁设备包含月租金（monthlyPrice）、设备总价（totalPrice）和租期（duration）信息，租期结束后设备所有权转移；经营租赁设备包含日租金（dailyPrice）、周租金（weeklyPrice）和月租金（monthlyPrice）信息，支持灵活的按需租赁模式。每个租赁设备还包含服务优势（benefits，JSON数组）、标签（tags，JSON数组）和评分（rating）等属性。')

add_heading2(doc, '6. 内容管理模块')

add_body_text(doc, '内容管理模块通过CmsContentController提供行业内容的发布与浏览功能，支持三种内容类型：video（视频）、article（文章）、vlog；四种内容分类：review（测评）、tutorial（教程）、vlog、news（资讯）。内容支持中英文双标题（title/titleEn），包含封面图（cover）、内容URL或正文（contentUrl）、作者（author）、标签（tags，JSON数组）等字段。内容详情接口自动递增浏览量。模块支持完整的CRUD操作。')

add_heading2(doc, '7. 订单与购物车模块')

add_body_text(doc, '该模块由购物车和订单两个子模块组成，是系统核心交易业务的实现载体。')

add_heading3(doc, '7.1 购物车子模块')
add_body_text(doc, '购物车服务（CartService）基于Redis Hash结构实现，以cart:user:{userId}为键，商品ID为Hash Field，CartItemDTO序列化对象为Hash Value。主要功能包括：')
cart_funcs = [
    '添加商品到购物车：查询商品信息后写入Redis Hash，若商品已存在则累加数量',
    '获取购物车列表：读取Redis Hash全部条目，反序列化为CartItemDTO列表',
    '更新商品数量：修改指定商品项的quantity字段',
    '删除单个商品：从Redis Hash中删除指定Field',
    '清空购物车：删除整个Redis Key',
    '购物车数据设置30天自动过期（EXPIRE命令）',
]
for item in cart_funcs:
    p = doc.add_paragraph(style='Normal')
    run = p.add_run(f'• {item}')
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(12)
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.5
    pf.left_indent = Cm(1.5)

add_heading3(doc, '7.2 订单子模块')
add_body_text(doc, '订单服务（OmsOrderService）支持两种下单方式：')
add_body_text(doc, '（1）购物车下单（createOrder）：从Redis购物车中读取数据，根据前端提交的商品ID列表（CreateOrderDTO.productIds）筛选选中的商品项，自动计算订单总金额（单价×数量求和），通过Hutool的Snowflake算法生成全局唯一订单号，在事务中依次创建订单主表（oms_order）和订单明细表（oms_order_item）记录，订单明细中保存商品名称、图片、价格等快照信息以防止商品信息变更影响历史订单。下单完成后自动清除购物车中已下单的商品项。')
add_body_text(doc, '（2）直接下单（createDirectOrder）：无需经过购物车，前端直接提交商品信息（DirectOrderDTO），系统据此创建订单主表和单条订单明细记录。')
add_body_text(doc, '订单支持状态流转：待付款(0) → 待发货(1) → 已发货(2) → 已完成(3)，或待付款(0) → 已取消(4)。模拟支付（payOrder）将状态从0更新为1并记录支付时间；取消订单（cancelOrder）仅允许待付款状态的订单执行。')

add_heading2(doc, '8. 用户交互模块')

add_body_text(doc, '用户交互模块包含四个子服务，分别对应四种用户交互行为的数据实体与业务逻辑：')

add_table(doc,
    headers=['子服务', '实体类', '数据表', '功能说明'],
    rows=[
        ['CommentService', 'PmsComment', 'pms_comment', '商品评论管理，支持1-5星评分、评论图片、有用数统计、状态控制（显示/隐藏）'],
        ['LikeService', 'UmsLike', 'ums_like', '点赞管理，支持多目标类型（product商品、comment评论、content内容），用户+目标唯一约束'],
        ['FavoriteService', 'UmsFavorite', 'ums_favorite', '收藏管理，支持多目标类型（product商品、supplier供应商、content内容），冗余存储目标名称和图片'],
        ['ShareService', 'UmsShare', 'ums_share', '分享记录管理，支持多平台（copy复制链接、wechat微信、weibo微博），允许未登录用户分享'],
    ],
    col_widths=[2.5, 2, 2, 9.5]
)

add_body_text(doc, '这些服务当前通过管理后台的AdminInteractionController进行数据的查询、状态更新、删除和批量操作。')

add_heading2(doc, '9. 管理后台模块')

add_body_text(doc, '管理后台模块提供系统运营所需的全面管理功能，由8个管理员控制器协同工作，全面覆盖用户管理、商品管理、供应商审核、订单管理、租赁设备管理和交互数据管理六大运营维度。')

admin_funcs = [
    ('AdminAuthController', '管理员专属登录接口，强制校验ADMIN角色，通过@PreAuthorize("hasRole(\'ADMIN\')")注解保护管理员信息查询接口。验证码校验为可选项。'),
    ('AdminDashboardController', '提供仪表盘统计数据，包括用户总数、供应商总数、待审核供应商数、商品总数、订单总数五项核心指标，以及最新注册用户列表和最新上架商品列表。'),
    ('AdminUserController', '用户完整CRUD操作、状态管理（封禁/解封）、密码重置（重置为默认密码123456）、批量删除等功能。内置管理员角色保护逻辑——不允许删除管理员账号、不允许修改管理员角色、不允许封禁管理员。'),
    ('AdminProductController', '商品CRUD、上架/下架控制、批量上下架与批量删除操作。创建商品时自动设置初始销量和浏览量为0。'),
    ('AdminSupplierController', '供应商CRUD、审核流程（通过时自动升级关联用户角色为SUPPLIER，拒绝时设置isVerified=2）、认证状态手动更新。审核操作在事务中执行。'),
    ('AdminOrderController', '订单列表查询（支持按订单号、状态筛选）、状态更新、删除（级联删除关联的订单明细记录）、批量删除，以及待处理订单数统计。'),
    ('AdminLeasingController', '租赁设备CRUD与上下架状态管理，创建设备时自动设置初始已租次数为0、评分为5.0。'),
    ('AdminInteractionController', '评论、点赞、收藏、分享四类交互数据的分页列表查询、状态更新、单条删除、批量删除，以及四类数据的总数统计。'),
]

for name, desc in admin_funcs:
    p = doc.add_paragraph(style='Normal')
    run_name = p.add_run(f'{name}：')
    run_name.font.name = '宋体'
    run_name._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run_name.font.size = Pt(12)
    run_name.font.bold = True
    run_desc = p.add_run(desc)
    run_desc.font.name = '宋体'
    run_desc._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run_desc.font.size = Pt(12)
    pf = p.paragraph_format
    pf.space_before = Pt(3)
    pf.space_after = Pt(3)
    pf.line_spacing = 1.5
    pf.first_line_indent = Cm(0.74)

doc.add_page_break()

# ============================================================
# 第四章 核心算法与流程
# ============================================================
add_heading1(doc, '第四章 核心算法与流程')

add_heading2(doc, '1. 用户认证与登录流程')

add_body_text(doc, '用户登录是系统的核心安全流程，涉及图形验证码校验、密码验证、角色匹配和JWT令牌生成等关键步骤。完整流程如下：')

login_steps = [
    '步骤1：客户端提交登录请求，携带用户名、密码、验证码UUID和验证码文本。',
    '步骤2：系统检查验证码UUID和Code是否为空，若为空则返回400错误"请输入验证码"。',
    '步骤3：从Redis中获取存储的验证码（Key格式：captcha:{uuid}）。若不存在，返回400错误"验证码已过期"。',
    '步骤4：立即删除Redis中的验证码记录（一次性消费），忽略大小写比较用户输入与存储值。若不匹配，返回400错误"验证码错误"。',
    '步骤5：根据用户名查询数据库（SysUserMapper.selectOne），若用户不存在则抛出CustomException"用户不存在"。',
    '步骤6：使用BCryptPasswordEncoder.matches()比对密码，若不匹配则抛出异常"密码错误"。',
    '步骤7：检查用户状态字段（status），若为0（禁用）则抛出异常"账号已禁用"。',
    '步骤8：执行角色校验逻辑——ADMIN角色可跨角色登录（直接通过）；非管理员角色需严格匹配数据库角色与登录角色参数。',
    '步骤9：调用JwtUtil.generateToken()，以username为Subject、签发时间为当前时间、过期时间为24小时后的JWT令牌，使用HS256算法签名。',
    '步骤10：将JWT令牌封装在Result.success()中返回客户端。',
]
for step in login_steps:
    p = doc.add_paragraph(style='Normal')
    run = p.add_run(step)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(12)
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.5
    pf.first_line_indent = Cm(0.74)

add_body_text(doc, '所有业务异常（CustomException）均由GlobalExceptionHandler统一捕获，封装为包含错误码和错误消息的Result响应返回客户端。')

add_heading2(doc, '2. JWT请求认证过滤流程')

add_body_text(doc, '每个HTTP请求到达控制器前，均经过JwtAuthenticationTokenFilter（继承OncePerRequestFilter）进行令牌解析与用户身份加载。该过滤器采用非阻断式设计，具体处理逻辑如下：')

jwt_steps = [
    '步骤1：从请求头获取Authorization字段值。',
    '步骤2：判断该值是否非空且以"Bearer "开头，若不满足则直接放行请求（filterChain.doFilter）。',
    '步骤3：截取Bearer后的Token字符串，调用JwtUtil.getClaimsByToken()解析JWT。',
    '步骤4：若解析抛出异常（Token无效或过期），在catch块中静默捕获，直接放行请求而不拦截。',
    '步骤5：从Claims中获取Subject（即username），检查username非空且SecurityContext中无已有认证信息。',
    '步骤6：根据username查询数据库获取SysUser对象，读取角色字段。',
    '步骤7：构造包含ROLE_{role}权限的UsernamePasswordAuthenticationToken对象，设入SecurityContextHolder。',
    '步骤8：放行请求，后续使用@PreAuthorize注解的接口可基于SecurityContext中的角色信息进行权限判断。',
]
for step in jwt_steps:
    p = doc.add_paragraph(style='Normal')
    run = p.add_run(step)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(12)
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.5
    pf.first_line_indent = Cm(0.74)

add_heading2(doc, '3. 用户注册流程')

add_body_text(doc, '用户注册流程涉及图形验证码、邮箱验证码双重验证以及密码加密等关键步骤，完整流程为：')

reg_steps = [
    '步骤1：客户端提交注册请求（RegisterDTO），包含用户名、密码、确认密码、邮箱、昵称、角色、验证码信息。',
    '步骤2：校验图形验证码——从Redis获取并一次性消费，忽略大小写比较。',
    '步骤3：校验邮箱验证码——从Redis获取（Key: email:code:register:{email}），一次性消费，精确比较。',
    '步骤4：检查password与confirmPassword是否一致。',
    '步骤5：构造SysUser对象，调用SysUserService.register()。',
    '步骤6：服务层检查用户名唯一性（通过LambdaQueryWrapper查询计数）。',
    '步骤7：使用BCryptPasswordEncoder.encode()对密码进行加密。',
    '步骤8：设置默认状态为启用（status=1）、默认角色为USER（若未指定）、设置创建时间和更新时间。',
    '步骤9：通过MyBatis-Plus的save()方法持久化至sys_user表，返回注册成功。',
]
for step in reg_steps:
    p = doc.add_paragraph(style='Normal')
    run = p.add_run(step)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(12)
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.5
    pf.first_line_indent = Cm(0.74)

add_heading2(doc, '4. 购物车与订单创建流程')

add_body_text(doc, '订单创建流程从购物车选品到订单生成，涉及Redis数据读取、金额计算、Snowflake唯一订单号生成和事务性数据库写入等核心处理逻辑。以购物车下单为例，完整流程为：')

order_steps = [
    '步骤1：用户提交下单请求（CreateOrderDTO），包含商品ID列表（productIds）等信息。',
    '步骤2：从Redis Hash（Key: cart:user:{userId}）中读取用户完整购物车数据。',
    '步骤3：将购物车中的CartItemDTO列表与前端提交的productIds进行交叉筛选，得到本次下单的商品项。',
    '步骤4：若筛选结果为空，抛出CustomException"请选择要购买的商品"。',
    '步骤5：计算订单总金额——将每件商品的单价（productPrice）乘以数量（quantity），使用BigDecimal.reduce求和。',
    '步骤6：调用Hutool的IdUtil.getSnowflakeNextIdStr()生成全局唯一的Snowflake订单号。',
    '步骤7：在@Transactional事务中，构造OmsOrder对象（订单号、用户ID、总金额、状态=0待付款、创建时间），通过save()方法写入oms_order表。',
    '步骤8：遍历选中的商品项，为每项构造OmsOrderItem对象（订单ID、商品ID、名称快照、图片快照、价格快照、数量、小计金额），通过OmsOrderItemMapper.insert()写入oms_order_item表。',
    '步骤9：调用CartService.removeItems()从Redis购物车中删除已下单的商品项。',
    '步骤10：事务提交，返回订单号。若任何步骤异常，事务整体回滚（rollbackFor = Exception.class）。',
]
for step in order_steps:
    p = doc.add_paragraph(style='Normal')
    run = p.add_run(step)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(12)
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.5
    pf.first_line_indent = Cm(0.74)

add_heading2(doc, '5. 订单状态流转')

add_body_text(doc, '订单在生命周期内经历多种状态，状态流转由用户操作和管理员操作共同驱动：')

add_table(doc,
    headers=['当前状态', '目标状态', '触发操作', '操作人', '附加行为'],
    rows=[
        ['待付款（0）', '待发货（1）', 'payOrder（模拟支付）', '用户', '记录支付时间（payTime）'],
        ['待付款（0）', '已取消（4）', 'cancelOrder（取消订单）', '用户', '仅待付款状态可取消'],
        ['待发货（1）', '已发货（2）', 'updateStatus（状态更新）', '管理员', '更新订单状态'],
        ['已发货（2）', '已完成（3）', 'updateStatus（状态更新）', '管理员', '更新订单状态'],
    ],
    col_widths=[2.5, 2.5, 3.5, 2, 5.5]
)
add_figure_caption(doc, '图4-1 订单状态流转表')

add_heading2(doc, '6. 供应商审核流程')

add_body_text(doc, '供应商审核是管理后台的核心业务流程之一，审核通过时会联动更新关联用户的角色权限。完整流程为：')

audit_steps = [
    '步骤1：管理员提交审核请求（AuditDTO），包含供应商ID和是否通过（pass）标志。',
    '步骤2：根据ID查询供应商记录，若不存在返回404错误。',
    '步骤3：若审核通过（pass=true）：更新供应商认证状态isVerified=1（已认证）；检查该供应商是否关联了userId，若有则通过LambdaUpdateWrapper将关联用户的角色更新为SUPPLIER。',
    '步骤4：若审核拒绝（pass=false）：更新供应商认证状态isVerified=2（拒绝）。',
    '步骤5：整个审核操作在@Transactional事务中执行，确保供应商状态更新和用户角色升级的原子性。',
]
for step in audit_steps:
    p = doc.add_paragraph(style='Normal')
    run = p.add_run(step)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(12)
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.5
    pf.first_line_indent = Cm(0.74)

doc.add_page_break()

# ============================================================
# 第五章 数据结构设计
# ============================================================
add_heading1(doc, '第五章 数据结构设计')

add_heading2(doc, '1. 核心数据模型')

add_body_text(doc, '系统共包含12张核心数据表，涵盖用户、供应商、商品、分类、订单、租赁、内容及用户交互等领域。所有数据表使用InnoDB引擎、utf8mb4字符集，主键采用BIGINT自增策略。以下详细列出每张表的字段结构：')

# sys_user
add_heading3(doc, '1.1 系统用户表（sys_user）')
add_table(doc,
    headers=['字段名', '类型', '约束', '说明'],
    rows=[
        ['id', 'BIGINT', '主键，自增', '主键ID'],
        ['username', 'VARCHAR(50)', 'NOT NULL，唯一索引', '用户名'],
        ['password', 'VARCHAR(100)', 'NOT NULL', '加密密码（BCrypt）'],
        ['nickname', 'VARCHAR(50)', '', '昵称'],
        ['email', 'VARCHAR(100)', '', '邮箱'],
        ['phone', 'VARCHAR(20)', '', '手机号'],
        ['avatar', 'VARCHAR(255)', '', '头像URL'],
        ['role', 'VARCHAR(20)', '默认USER', '角色：USER/SUPPLIER/ADMIN'],
        ['status', 'TINYINT', '默认1', '状态：1-正常，0-禁用'],
        ['create_time', 'DATETIME', '默认当前时间', '创建时间'],
        ['update_time', 'DATETIME', '自动更新', '更新时间'],
    ],
    col_widths=[3, 3, 3.5, 6.5]
)

# pms_product
add_heading3(doc, '1.2 商品表（pms_product）')
add_table(doc,
    headers=['字段名', '类型', '约束', '说明'],
    rows=[
        ['id', 'BIGINT', '主键，自增', '主键ID'],
        ['supplier_id', 'BIGINT', 'NOT NULL', '供应商ID'],
        ['category_id', 'BIGINT', 'NOT NULL', '分类ID'],
        ['name', 'VARCHAR(200)', 'NOT NULL', '商品名称'],
        ['price', 'DECIMAL(10,2)', 'NOT NULL', '售价'],
        ['original_price', 'DECIMAL(10,2)', '', '原价'],
        ['stock', 'INT', '默认0', '库存'],
        ['image', 'VARCHAR(255)', '', '主图URL'],
        ['album', 'TEXT', '', '相册图片（JSON数组）'],
        ['description', 'TEXT', '', '商品描述'],
        ['specs', 'TEXT', '', '规格参数（JSON格式）'],
        ['status', 'TINYINT', '默认1', '状态：1-上架，0-下架'],
        ['sales', 'INT', '默认0', '销量'],
        ['views', 'INT', '默认0', '浏览量'],
        ['create_time', 'DATETIME', '默认当前时间', '创建时间'],
    ],
    col_widths=[3, 3, 3.5, 6.5]
)

# oms_order
add_heading3(doc, '1.3 订单表（oms_order）')
add_table(doc,
    headers=['字段名', '类型', '约束', '说明'],
    rows=[
        ['id', 'BIGINT', '主键，自增', '主键ID'],
        ['order_no', 'VARCHAR(64)', 'NOT NULL，唯一索引', '订单号（Snowflake）'],
        ['user_id', 'BIGINT', 'NOT NULL', '用户ID'],
        ['total_amount', 'DECIMAL(10,2)', 'NOT NULL', '订单总金额'],
        ['status', 'TINYINT', '默认0', '0-待付款/1-待发货/2-已发货/3-已完成/4-已取消'],
        ['pay_time', 'DATETIME', '', '支付时间'],
        ['delivery_time', 'DATETIME', '', '发货时间'],
        ['finish_time', 'DATETIME', '', '完成时间'],
        ['create_time', 'DATETIME', '默认当前时间', '创建时间'],
        ['update_time', 'DATETIME', '', '更新时间'],
    ],
    col_widths=[3, 3, 3.5, 6.5]
)

# oms_order_item
add_heading3(doc, '1.4 订单明细表（oms_order_item）')
add_table(doc,
    headers=['字段名', '类型', '约束', '说明'],
    rows=[
        ['id', 'BIGINT', '主键，自增', '主键ID'],
        ['order_id', 'BIGINT', 'NOT NULL，索引', '订单ID'],
        ['product_id', 'BIGINT', 'NOT NULL', '商品ID'],
        ['product_name', 'VARCHAR(200)', '', '商品名称快照'],
        ['product_image', 'VARCHAR(255)', '', '商品图片快照'],
        ['product_price', 'DECIMAL(10,2)', '', '商品价格快照'],
        ['quantity', 'INT', '默认1', '购买数量'],
        ['subtotal', 'DECIMAL(10,2)', '', '小计金额'],
    ],
    col_widths=[3, 3, 3.5, 6.5]
)

# oms_supplier
add_heading3(doc, '1.5 供应商表（oms_supplier）')
add_table(doc,
    headers=['字段名', '类型', '约束', '说明'],
    rows=[
        ['id', 'BIGINT', '主键，自增', '主键ID'],
        ['name', 'VARCHAR(100)', 'NOT NULL', '供应商名称'],
        ['logo', 'VARCHAR(255)', '', 'Logo图片URL'],
        ['description', 'TEXT', '', '企业简介'],
        ['contact_info', 'VARCHAR(500)', '', '联系方式（JSON格式）'],
        ['is_verified', 'TINYINT', '默认0', '0-待审核/1-已认证/2-拒绝'],
        ['create_time', 'DATETIME', '默认当前时间', '创建时间'],
        ['user_id', 'BIGINT', '', '关联用户ID'],
    ],
    col_widths=[3, 3, 3.5, 6.5]
)

# pms_category
add_heading3(doc, '1.6 商品分类表（pms_category）')
add_table(doc,
    headers=['字段名', '类型', '约束', '说明'],
    rows=[
        ['id', 'BIGINT', '主键，自增', '主键ID'],
        ['name', 'VARCHAR(50)', 'NOT NULL', '分类名称'],
        ['parent_id', 'BIGINT', '默认0', '父分类ID，0为顶级'],
        ['sort', 'INT', '默认0', '排序值'],
        ['icon', 'VARCHAR(255)', '', '图标'],
    ],
    col_widths=[3, 3, 3.5, 6.5]
)

# oms_leasing
add_heading3(doc, '1.7 租赁设备表（oms_leasing）')
add_table(doc,
    headers=['字段名', '类型', '说明'],
    rows=[
        ['id', 'BIGINT', '主键ID'],
        ['name', 'VARCHAR(200)', '设备名称'],
        ['type', 'VARCHAR(20)', '融资租赁financing/经营租赁operating'],
        ['image', 'VARCHAR(255)', '设备图片URL'],
        ['description', 'TEXT', '设备描述'],
        ['supplier', 'VARCHAR(100)', '供应商名称'],
        ['supplier_id', 'BIGINT', '供应商ID'],
        ['monthly_price', 'DECIMAL(10,2)', '月租金'],
        ['total_price', 'DECIMAL(10,2)', '设备总价（融资租赁）'],
        ['duration', 'VARCHAR(50)', '租期'],
        ['daily_price', 'DECIMAL(10,2)', '日租金（经营租赁）'],
        ['weekly_price', 'DECIMAL(10,2)', '周租金（经营租赁）'],
        ['benefits', 'TEXT', '服务优势（JSON数组）'],
        ['tags', 'TEXT', '标签（JSON数组）'],
        ['leased', 'INT', '已租次数'],
        ['rating', 'DECIMAL(2,1)', '评分（默认5.0）'],
        ['status', 'TINYINT', '1-上架/0-下架'],
        ['create_time', 'DATETIME', '创建时间'],
        ['update_time', 'DATETIME', '更新时间'],
    ],
    col_widths=[3, 3, 10]
)

# 其余表汇总
add_heading3(doc, '1.8 其他数据表汇总')
add_table(doc,
    headers=['表名', '实体类', '主要字段', '关键约束'],
    rows=[
        ['cms_content', 'CmsContent', 'title, titleEn, type, cover, contentUrl, author, views, category, tags, status', '无外键约束'],
        ['pms_comment', 'PmsComment', 'productId, userId, userName, userAvatar, rating, content, images, helpful, status', 'product_id索引，user_id索引'],
        ['ums_like', 'UmsLike', 'userId, targetType, targetId', 'userId+targetType+targetId唯一索引'],
        ['ums_favorite', 'UmsFavorite', 'userId, targetType, targetId, targetName, targetImage', 'userId+targetType+targetId唯一索引'],
        ['ums_share', 'UmsShare', 'userId, targetType, targetId, targetName, shareUrl, platform', 'targetType+targetId索引'],
    ],
    col_widths=[2.5, 2.5, 6.5, 4.5]
)

add_heading2(doc, '2. 数据实体间关系')

add_body_text(doc, '系统数据模型以SysUser为核心用户实体，通过外键关联辐射至供应商、订单、评论、点赞、收藏和分享等业务实体。核心关联关系如下：')

add_table(doc,
    headers=['关系', '说明', '类型'],
    rows=[
        ['SysUser → OmsSupplier', '用户注册为供应商后关联', '一对多（userId）'],
        ['SysUser → OmsOrder', '用户下单', '一对多（userId）'],
        ['SysUser → PmsComment/UmsLike/UmsFavorite/UmsShare', '用户交互行为', '一对多（userId）'],
        ['OmsSupplier → PmsProduct', '供应商供货', '一对多（supplierId）'],
        ['OmsSupplier → OmsLeasing', '供应商出租', '一对多（supplierId）'],
        ['PmsCategory → PmsProduct', '商品分类', '一对多（categoryId）'],
        ['PmsProduct → PmsComment', '商品评论', '一对多（productId）'],
        ['PmsProduct → OmsOrderItem', '商品订购', '一对多（productId）'],
        ['OmsOrder → OmsOrderItem', '订单包含明细', '一对多（orderId）'],
    ],
    col_widths=[5, 5, 6]
)
add_figure_caption(doc, '图5-1 数据实体关联关系表')

add_heading2(doc, '3. 缓存数据结构')

add_body_text(doc, '系统使用Redis存储以下三类临时数据，用于提升性能和支撑特定业务场景：')

add_table(doc,
    headers=['缓存Key格式', '数据类型', '有效期', '用途说明'],
    rows=[
        ['captcha:{uuid}', 'String', '5分钟', '图形验证码文本，一次性消费'],
        ['email:code:{type}:{email}', 'String', '5分钟', '邮箱验证码（type: register/reset/login）'],
        ['cart:user:{userId}', 'Hash', '30天', '购物车数据，Field=商品ID，Value=CartItemDTO'],
    ],
    col_widths=[4.5, 2.5, 2, 7]
)

add_heading2(doc, '4. 数据传输对象（DTO）')

add_body_text(doc, '系统定义了以下核心数据传输对象，用于控制器层与客户端之间的数据交换：')

add_table(doc,
    headers=['DTO名称', '所属模块', '主要字段', '用途'],
    rows=[
        ['LoginDTO', 'system', 'username, password, captchaUuid, captchaCode, role', '用户登录请求'],
        ['RegisterDTO', 'system', 'username, password, confirmPassword, email, nickname, role, captchaUuid, captchaCode, emailCode, phone', '用户注册请求'],
        ['ResetPasswordDTO', 'system', 'username, email, emailCode, newPassword, confirmPassword, captchaUuid, captchaCode', '密码重置请求'],
        ['AdminLoginDTO', 'admin', 'username, password, captchaUuid, captchaCode', '管理员登录请求'],
        ['AuditDTO', 'admin', 'id, pass', '供应商审核请求'],
        ['UserStatusDTO', 'admin', 'userId, status', '用户状态变更请求'],
        ['CartItemDTO', 'order', 'productId, productName, productImage, productPrice, quantity, checked', '购物车项（Redis存储）'],
        ['CreateOrderDTO', 'order', 'productIds, address, receiverName, receiverPhone, remark', '购物车下单请求'],
        ['DirectOrderDTO', 'order', 'productId, productName, productImage, productPrice, quantity, address, receiverName, receiverPhone, remark', '直接下单请求'],
    ],
    col_widths=[2.5, 1.5, 7, 5]
)

doc.add_page_break()

# ============================================================
# 第六章 接口设计
# ============================================================
add_heading1(doc, '第六章 接口设计')

add_heading2(doc, '1. 接口总体说明')

add_body_text(doc, '本系统所有API接口遵循RESTful设计规范，采用JSON作为请求与响应的数据格式。所有接口响应均通过统一的Result<T>包装类返回，结构如下：')

add_table(doc,
    headers=['字段', '类型', '说明'],
    rows=[
        ['code', 'Integer', '状态码。200-成功，400-参数错误，401-未认证，403-无权限，404-资源不存在，500-服务器错误'],
        ['message', 'String', '响应消息。成功时为"success"，失败时为具体错误描述'],
        ['data', 'T（泛型）', '响应数据。类型根据接口不同而变化，失败时为null'],
    ],
    col_widths=[2, 3, 11]
)

add_body_text(doc, '需认证的接口通过HTTP请求头携带JWT令牌：Authorization: Bearer {token}。')

add_heading2(doc, '2. 接口模块划分')

add_body_text(doc, '系统接口分为用户端接口和管理端接口两大组，共计10个接口模块：')

add_table(doc,
    headers=['分组', '接口前缀', '模块说明'],
    rows=[
        ['用户端', '/api/v1/user', '登录、注册、重置密码、用户信息获取'],
        ['用户端', '/api/v1/captcha', '获取图形验证码'],
        ['用户端', '/api/v1/email', '发送邮箱验证码'],
        ['用户端', '/api/product', '商品列表、详情、分类'],
        ['用户端', '/api/v1/category', '分类CRUD管理'],
        ['用户端', '/api/v1/supplier', '供应商列表、详情、CRUD'],
        ['用户端', '/api/leasing', '租赁设备列表、详情'],
        ['用户端', '/api/v1/content', '内容列表、详情、CRUD'],
        ['用户端', '/api/v1/cart', '购物车增删改查、清空'],
        ['用户端', '/api/v1/order', '创建订单、列表、详情、支付、取消'],
        ['管理端', '/api/admin', '管理员登录、信息获取'],
        ['管理端', '/api/admin/dashboard', '统计数据、最新用户、最新商品'],
        ['管理端', '/api/admin/user', '用户CRUD、封禁、重置密码、批量操作'],
        ['管理端', '/api/admin/product', '商品CRUD、上下架、批量操作'],
        ['管理端', '/api/admin/supplier', '供应商CRUD、审核、认证状态'],
        ['管理端', '/api/admin/order', '订单查询、状态更新、删除、统计'],
        ['管理端', '/api/admin/leasing', '租赁设备CRUD、上下架'],
        ['管理端', '/api/admin/interaction', '评论/点赞/收藏/分享管理'],
    ],
    col_widths=[2, 4, 10]
)
add_figure_caption(doc, '图6-1 系统接口模块总览')

add_heading2(doc, '3. 系统认证接口')
add_table(doc,
    headers=['端点', '方法', '说明', '主要参数', '返回值'],
    rows=[
        ['/api/v1/user/login', 'POST', '用户登录', 'LoginDTO', 'Result<String>（JWT令牌）'],
        ['/api/v1/user/register', 'POST', '用户注册', 'RegisterDTO', 'Result<String>'],
        ['/api/v1/user/reset-password', 'POST', '密码重置', 'ResetPasswordDTO', 'Result<String>'],
        ['/api/v1/user/info', 'GET', '获取用户信息', 'Principal（JWT解析）', 'Result<SysUser>'],
        ['/api/v1/captcha/image', 'GET', '获取图形验证码', '无', 'Result<Map>（uuid, imageBase64）'],
        ['/api/v1/email/send', 'POST', '发送邮箱验证码', 'Map（email, type）', 'Result<String>'],
    ],
    col_widths=[3.5, 1.2, 2.5, 4.3, 4.5]
)

add_heading2(doc, '4. 商品与分类接口')
add_table(doc,
    headers=['端点', '方法', '说明', '主要参数', '返回值'],
    rows=[
        ['/api/product/list', 'GET', '商品列表', 'page, size, name, categoryId, sort, includeOffShelf', 'Result<IPage<PmsProduct>>'],
        ['/api/product/{id}', 'GET', '商品详情', 'id（路径参数）', 'Result<PmsProduct>'],
        ['/api/product/categories', 'GET', '获取分类列表', '无', 'Result<List<PmsCategory>>'],
        ['/api/v1/category/list', 'GET', '分类列表', '无', 'Result<List<PmsCategory>>'],
        ['/api/v1/category', 'POST', '创建分类', 'PmsCategory', 'Result<String>'],
        ['/api/v1/category', 'PUT', '更新分类', 'PmsCategory', 'Result<String>'],
        ['/api/v1/category/{id}', 'DELETE', '删除分类', 'id', 'Result<String>'],
    ],
    col_widths=[3.5, 1.2, 2.5, 4.3, 4.5]
)

add_heading2(doc, '5. 购物车与订单接口')
add_table(doc,
    headers=['端点', '方法', '说明', '主要参数', '返回值'],
    rows=[
        ['/api/v1/cart', 'GET', '获取购物车', 'Principal', 'Result<List<CartItemDTO>>'],
        ['/api/v1/cart/add', 'POST', '加入购物车', 'productId, quantity', 'Result<String>'],
        ['/api/v1/cart/update', 'PUT', '更新数量', 'productId, quantity', 'Result<String>'],
        ['/api/v1/cart/remove/{id}', 'DELETE', '移除商品', 'productId', 'Result<String>'],
        ['/api/v1/cart/clear', 'DELETE', '清空购物车', 'Principal', 'Result<String>'],
        ['/api/v1/order/create', 'POST', '购物车下单', 'CreateOrderDTO', 'Result<String>（订单号）'],
        ['/api/v1/order/direct', 'POST', '直接下单', 'DirectOrderDTO', 'Result<String>（订单号）'],
        ['/api/v1/order/list', 'GET', '用户订单列表', 'page, size, status', 'Result<Page<OmsOrder>>'],
        ['/api/v1/order/{orderId}', 'GET', '订单详情', 'orderId', 'Result<Map>（order, items）'],
        ['/api/v1/order/pay/{id}', 'POST', '模拟支付', 'orderId', 'Result<String>'],
        ['/api/v1/order/cancel/{id}', 'POST', '取消订单', 'orderId', 'Result<String>'],
    ],
    col_widths=[3.5, 1.2, 2.5, 4.3, 4.5]
)

add_heading2(doc, '6. 管理后台核心接口')
add_table(doc,
    headers=['端点', '方法', '说明'],
    rows=[
        ['/api/admin/login', 'POST', '管理员登录（强制ADMIN角色校验）'],
        ['/api/admin/info', 'GET', '获取管理员信息（@PreAuthorize保护）'],
        ['/api/admin/dashboard/stats', 'GET', '仪表盘统计（用户/供应商/商品/订单计数）'],
        ['/api/admin/dashboard/recent-users', 'GET', '最新注册用户列表'],
        ['/api/admin/dashboard/recent-products', 'GET', '最新上架商品列表'],
        ['/api/admin/supplier/audit', 'POST', '供应商审核（通过/拒绝）'],
        ['/api/admin/supplier/status', 'POST', '更新供应商认证状态'],
        ['/api/admin/product/off-shelf', 'POST', '强制下架商品'],
        ['/api/admin/product/on-shelf', 'POST', '上架商品'],
        ['/api/admin/product/batch/on-shelf', 'POST', '批量上架'],
        ['/api/admin/product/batch/off-shelf', 'POST', '批量下架'],
        ['/api/admin/order/{id}/status', 'PUT', '更新订单状态'],
        ['/api/admin/user/{id}/status', 'PUT', '封禁/解封用户'],
        ['/api/admin/user/{id}/reset-password', 'PUT', '重置用户密码'],
        ['/api/admin/interaction/comment/list', 'GET', '评论列表'],
        ['/api/admin/interaction/like/list', 'GET', '点赞列表'],
        ['/api/admin/interaction/stats', 'GET', '交互数据统计'],
    ],
    col_widths=[5, 1.5, 9.5]
)

add_heading2(doc, '7. 接口调用时序说明')

add_body_text(doc, '以下描述用户从登录到下单的典型接口调用时序：')

seq_steps = [
    '1. 客户端 → 后端：GET /api/v1/captcha/image【获取验证码】',
    '   后端 → Redis：存储验证码（captcha:{uuid}, 5分钟有效）',
    '   后端 → 客户端：返回uuid + imageBase64数据',
    '',
    '2. 客户端 → 后端：POST /api/v1/user/login（LoginDTO）【用户登录】',
    '   后端 → Redis：获取并删除验证码',
    '   后端 → MySQL：查询用户记录，校验密码和角色',
    '   后端 → 客户端：返回JWT令牌',
    '',
    '3. 客户端 → 后端：GET /api/product/list（携带Authorization头）【浏览商品】',
    '   后端 → MySQL：分页查询商品列表',
    '   后端 → 客户端：返回商品分页数据',
    '',
    '4. 客户端 → 后端：POST /api/v1/cart/add（productId, quantity）【加入购物车】',
    '   后端 → MySQL：查询商品信息',
    '   后端 → Redis：写入购物车Hash（cart:user:{userId}）',
    '   后端 → 客户端：返回成功',
    '',
    '5. 客户端 → 后端：POST /api/v1/order/create（CreateOrderDTO）【提交订单】',
    '   后端 → Redis：读取购物车数据',
    '   后端 → MySQL：事务写入oms_order + oms_order_item',
    '   后端 → Redis：删除已下单的购物车项',
    '   后端 → 客户端：返回订单号',
    '',
    '6. 客户端 → 后端：POST /api/v1/order/pay/{orderId}【模拟支付】',
    '   后端 → MySQL：更新订单状态为待发货',
    '   后端 → 客户端：返回支付成功',
]

for step in seq_steps:
    if step == '':
        doc.add_paragraph()
    else:
        p = doc.add_paragraph(style='Normal')
        run = p.add_run(step)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(11)
        pf = p.paragraph_format
        pf.space_before = Pt(1)
        pf.space_after = Pt(1)
        pf.line_spacing = 1.5
        pf.left_indent = Cm(1)

add_figure_caption(doc, '图6-2 用户下单典型调用时序')

doc.add_page_break()

# ============================================================
# 第七章 异常处理设计
# ============================================================
add_heading1(doc, '第七章 异常处理设计')

add_heading2(doc, '1. 全局异常处理机制')

add_body_text(doc, '本系统通过GlobalExceptionHandler类（标注@RestControllerAdvice注解）实现全局异常的统一捕获与响应封装。异常处理分为两个层级：')

add_body_text(doc, '（1）业务异常（CustomException）：由业务逻辑中主动抛出，继承RuntimeException，携带自定义错误码（code，Integer类型）和错误消息（message，String类型）。例如"用户不存在"、"密码错误"、"验证码已过期"等可预见的业务校验失败场景。处理方法：handleCustomException()，日志级别ERROR，返回Result.error(code, message)。')

add_body_text(doc, '（2）系统异常（Exception）：捕获所有未被业务异常处理器拦截的运行时异常和受检异常，统一返回500错误码并附带异常信息"System Error: {异常消息}"。处理方法：handleException()，日志级别ERROR并输出堆栈信息。')

add_body_text(doc, '此外，JWT过滤器（JwtAuthenticationTokenFilter）中的令牌解析异常采用静默捕获策略——在try-catch块中捕获所有异常但不做任何处理，确保令牌问题不会阻断正常请求的处理流程。')

add_heading2(doc, '2. 异常处理流程')

add_body_text(doc, '全局异常处理的完整流程如下：')

exc_steps = [
    '步骤1：Controller或Service层在执行业务逻辑过程中发生异常。',
    '步骤2：若为主动抛出的CustomException，由handleCustomException方法捕获。',
    '步骤3：记录ERROR日志"Business Exception: {message}"。',
    '步骤4：返回Result.error(自定义code, message)封装的JSON响应。',
    '步骤5：若为其他未预见的Exception，由handleException方法捕获。',
    '步骤6：记录ERROR日志"System Exception: {message}"并输出完整堆栈跟踪。',
    '步骤7：返回Result.error(500, "System Error: " + message)封装的JSON响应。',
    '步骤8：JWT过滤器中的异常单独处理——catch块静默吞没异常，请求继续放行。',
]
for step in exc_steps:
    p = doc.add_paragraph(style='Normal')
    run = p.add_run(step)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(12)
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.5
    pf.first_line_indent = Cm(0.74)

add_heading2(doc, '3. 主要异常场景')

add_body_text(doc, '系统代码中已显式处理的主要异常场景统计如下：')

add_table(doc,
    headers=['异常场景', '处理位置', '错误码', '错误消息'],
    rows=[
        ['图形验证码为空', 'SysUserController', '400', '请输入验证码'],
        ['图形验证码错误或过期', 'SysUserController', '400', '验证码错误或已过期'],
        ['邮箱验证码为空', 'SysUserController', '400', '请输入邮箱验证码'],
        ['邮箱验证码错误或过期', 'SysUserController', '400', '邮箱验证码错误或已过期'],
        ['密码确认不一致', 'SysUserController', '400', '两次密码输入不一致'],
        ['用户不存在', 'SysUserService', '500', '用户不存在'],
        ['密码错误', 'SysUserService', '500', '密码错误'],
        ['账号已禁用', 'SysUserService', '500', '账号已禁用'],
        ['角色不匹配（供应商）', 'SysUserService', '500', '该账号不是供应商账号'],
        ['角色不匹配（采购方）', 'SysUserService', '500', '该账号不是采购方账号'],
        ['用户名已存在', 'SysUserService', '500', 'Username already exists'],
        ['用户名/邮箱不匹配', 'SysUserService', '500', '用户名或邮箱错误'],
        ['非管理员登录后台', 'AdminAuthController', '500', '无权访问管理后台'],
        ['商品不存在', 'PmsProductController', '404', '商品不存在'],
        ['供应商不存在', 'AdminSupplierController', '404', '供应商不存在'],
        ['设备不存在', 'OmsLeasingController', '404', '设备不存在'],
        ['订单不存在', 'OmsOrderService', '500', '订单不存在'],
        ['订单状态异常', 'OmsOrderService', '500', '订单状态异常'],
        ['仅待付款可取消', 'OmsOrderService', '500', '只有待付款订单可以取消'],
        ['购物车无选中商品', 'OmsOrderService', '500', '请选择要购买的商品'],
        ['删除管理员禁止', 'AdminUserController', '403', '不能删除管理员账号'],
        ['修改管理员角色禁止', 'AdminUserController', '403', '不能修改管理员角色'],
        ['封禁管理员禁止', 'AdminUserController', '403', '无法更改管理员状态'],
        ['邮箱格式校验失败', 'EmailController', '400', '邮箱格式不正确'],
        ['邮件发送失败', 'EmailService', '500', '验证码发送失败'],
        ['JWT令牌无效/过期', 'JwtAuthFilter', '—', '静默忽略，请求放行'],
        ['未登录访问', 'SysUserController', '401', '未登录'],
    ],
    col_widths=[3.5, 3.5, 1.5, 7.5]
)
add_figure_caption(doc, '图7-1 系统异常场景清单')

add_heading2(doc, '4. 事务回滚策略')

add_body_text(doc, '系统在以下关键操作中启用了Spring声明式事务管理，通过@Transactional(rollbackFor = Exception.class)注解确保数据一致性：')

add_body_text(doc, '（1）订单创建（OmsOrderService.createOrder / createDirectOrder）：确保订单主表（oms_order）写入、订单明细表（oms_order_item）写入和购物车清理操作的原子性。任何步骤发生异常时，所有数据库操作全部回滚，购物车数据不受影响。')

add_body_text(doc, '（2）订单状态变更（OmsOrderService.payOrder / cancelOrder）：在事务保护下更新订单状态和时间戳字段，避免并发场景下的状态不一致问题。')

add_body_text(doc, '（3）供应商审核（AdminSupplierController.audit）：确保供应商认证状态更新（isVerified字段）和关联用户角色升级（role字段更新为SUPPLIER）两个数据库操作的原子性，避免出现供应商已认证但用户角色未更新的不一致状态。')

# ============================
# 结束语
# ============================
doc.add_page_break()
add_formatted_paragraph(
    doc, '内容到此结束，感谢您的查看！',
    font_name='宋体', font_size=12, bold=False,
    alignment=WD_ALIGN_PARAGRAPH.CENTER,
    space_before=200, space_after=0
)

# ============================
# 保存文档
# ============================
output_path = r'F:\Development\Java\IDEA_Projects\MinJue\B2B工业设备宣传与电商平台_软件设计说明书_V1.0.docx'
doc.save(output_path)
print(f'文档已生成：{output_path}')
print(f'文件大小：{os.path.getsize(output_path)} bytes')
