from docx import Document
from docx.shared import Pt, Emu, Cm

doc = Document(r'F:\Development\Java\IDEA_Projects\MinJue\基于深度学习的智能类案推送平台使用手册.docx')

# Detailed font/style analysis
for i, p in enumerate(doc.paragraphs[:30]):
    text = p.text.strip()
    if not text and p.style.name == 'Normal':
        continue
    style_name = p.style.name
    align = p.alignment
    fonts = []
    for run in p.runs:
        f = run.font
        fonts.append({
            'text': run.text[:20],
            'name': f.name,
            'size': str(f.size) if f.size else None,
            'bold': f.bold,
            'color': str(f.color.rgb) if f.color and f.color.rgb else None,
        })
    pf = p.paragraph_format
    print(f'[{i}] Style={style_name} Align={align}')
    print(f'  Text: {text[:80]}')
    print(f'  SpaceBefore={pf.space_before} SpaceAfter={pf.space_after} LineSpacing={pf.line_spacing}')
    if fonts:
        for fi, font in enumerate(fonts):
            print(f'  Run{fi}: name={font["name"]} size={font["size"]} bold={font["bold"]} color={font["color"]} text="{font["text"]}"')
    print()

# Check heading styles
print('=== HEADING STYLES ===')
for p in doc.paragraphs:
    if p.style.name.startswith('Heading'):
        pf = p.paragraph_format
        fonts = []
        for run in p.runs:
            f = run.font
            fonts.append({'name': f.name, 'size': str(f.size) if f.size else None, 'bold': f.bold})
        print(f'Style={p.style.name} Text={p.text.strip()[:50]}')
        if fonts:
            print(f'  Fonts: {fonts}')
        print(f'  SpaceBefore={pf.space_before} SpaceAfter={pf.space_after}')
